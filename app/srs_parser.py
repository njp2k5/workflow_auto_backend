"""
SRS (Software Requirements Specification) Document Parser.
Rule-based extraction of sections from Word documents.
Maps SRS sections to Confluence pages and generates Jira tasks.
"""
import re
from typing import Dict, List, Optional, Any, TYPE_CHECKING
from dataclasses import dataclass, field
from io import BytesIO

from docx import Document

if TYPE_CHECKING:
    from docx.document import Document as DocxDocument

from app.logger import get_logger

logger = get_logger(__name__)


# Section mapping from SRS to Confluence pages
SRS_SECTION_MAPPING = {
    "introduction": "Product Overview",
    "scope": "System Scope",
    "user types": "Personas",
    "user type": "Personas",
    "personas": "Personas",
    "functional requirements": "Feature Pages",
    "functional": "Feature Pages",
    "features": "Feature Pages",
    "non-functional requirements": "NFR",
    "non-functional": "NFR",
    "nfr": "NFR",
    "performance": "NFR",
    "security": "NFR",
    "ui": "UI/UX",
    "user interface": "UI/UX",
    "ux": "UI/UX",
    "design": "UI/UX",
    "api": "API Docs",
    "api specification": "API Docs",
    "api endpoints": "API Docs",
    "interfaces": "API Docs",
    "workflows": "Diagrams",
    "workflow": "Diagrams",
    "diagrams": "Diagrams",
    "flow": "Diagrams",
    "process flow": "Diagrams",
}

# Team members for task assignment
TEAM_MEMBERS = [
    "Nikhil J Prasad",
    "S Govind Krishnan",
    "Kailas S S",
    "Mukundan V S",
]

# Task templates based on section type
TASK_TEMPLATES = {
    "Product Overview": [
        {"title": "Review and finalize product vision document", "assignee_idx": 0},
        {"title": "Define success metrics and KPIs", "assignee_idx": 1},
    ],
    "System Scope": [
        {"title": "Document system boundaries and constraints", "assignee_idx": 2},
        {"title": "Identify external system dependencies", "assignee_idx": 3},
    ],
    "Personas": [
        {"title": "Create detailed user persona profiles", "assignee_idx": 0},
        {"title": "Map user journeys for each persona", "assignee_idx": 1},
    ],
    "Feature Pages": [
        {"title": "Break down features into user stories", "assignee_idx": 0},
        {"title": "Define acceptance criteria for features", "assignee_idx": 1},
        {"title": "Estimate feature complexity and effort", "assignee_idx": 2},
        {"title": "Prioritize features for MVP", "assignee_idx": 3},
    ],
    "NFR": [
        {"title": "Define performance benchmarks and SLAs", "assignee_idx": 2},
        {"title": "Document security requirements and compliance", "assignee_idx": 3},
    ],
    "UI/UX": [
        {"title": "Create wireframes for main screens", "assignee_idx": 1},
        {"title": "Design component library specifications", "assignee_idx": 0},
    ],
    "API Docs": [
        {"title": "Document API endpoints and contracts", "assignee_idx": 2},
        {"title": "Define API authentication and authorization", "assignee_idx": 3},
    ],
    "Diagrams": [
        {"title": "Create system architecture diagram", "assignee_idx": 2},
        {"title": "Document workflow state diagrams", "assignee_idx": 3},
    ],
}


@dataclass
class SRSSection:
    """Represents a parsed section from the SRS document."""
    title: str
    content: str
    confluence_page: str
    subsections: List[Dict[str, str]] = field(default_factory=list)
    requirements: List[str] = field(default_factory=list)


@dataclass
class ParsedSRS:
    """Complete parsed SRS document."""
    document_title: str
    sections: List[SRSSection]
    raw_text: str
    metadata: Dict[str, Any] = field(default_factory=dict)


class SRSParser:
    """
    Rule-based parser for SRS documents in Word format.
    Extracts sections and maps them to Confluence page types.
    """
    
    def __init__(self):
        self.section_mapping = SRS_SECTION_MAPPING
        self.team_members = TEAM_MEMBERS
    
    def parse_document(self, file_content: bytes) -> ParsedSRS:
        """
        Parse a Word document and extract SRS sections.
        
        Args:
            file_content: Raw bytes of the Word document
            
        Returns:
            ParsedSRS object with extracted sections
        """
        logger.info("Starting SRS document parsing...")
        
        # Load document from bytes
        doc = Document(BytesIO(file_content))
        
        # Extract document title
        doc_title = self._extract_document_title(doc)
        logger.info(f"Document title: {doc_title}")
        
        # Extract all text for raw reference
        raw_text = self._get_full_text(doc)
        
        # Extract sections based on headings
        sections = self._extract_sections(doc)
        logger.info(f"Extracted {len(sections)} sections")
        
        # Extract metadata
        metadata = self._extract_metadata(doc)
        
        return ParsedSRS(
            document_title=doc_title,
            sections=sections,
            raw_text=raw_text,
            metadata=metadata
        )
    
    def _extract_document_title(self, doc: Document) -> str:
        """Extract the document title from the first heading or paragraph."""
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Check if it's a heading style
                if para.style.name.startswith('Heading') or para.style.name == 'Title':
                    return text
                # Otherwise use first non-empty paragraph
                return text
        return "Software Requirements Specification"
    
    def _get_full_text(self, doc: Document) -> str:
        """Get all text from the document."""
        return "\n".join([para.text for para in doc.paragraphs])
    
    def _extract_sections(self, doc: Document) -> List[SRSSection]:
        """
        Extract sections based on heading styles and content patterns.
        Uses rule-based matching to identify SRS sections.
        """
        sections = []
        current_section = None
        current_content = []
        current_requirements = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Check if this is a section heading
            is_heading = (
                para.style.name.startswith('Heading') or
                self._is_section_header(text)
            )
            
            if is_heading:
                # Save previous section if exists
                if current_section:
                    content = "\n".join(current_content)
                    current_section.content = content
                    current_section.requirements = current_requirements
                    sections.append(current_section)
                
                # Start new section
                confluence_page = self._map_to_confluence_page(text)
                current_section = SRSSection(
                    title=text,
                    content="",
                    confluence_page=confluence_page
                )
                current_content = []
                current_requirements = []
            else:
                if current_section:
                    current_content.append(text)
                    
                    # Extract requirements (lines starting with REQ, FR, NFR, etc.)
                    if self._is_requirement(text):
                        current_requirements.append(text)
        
        # Don't forget the last section
        if current_section:
            current_section.content = "\n".join(current_content)
            current_section.requirements = current_requirements
            sections.append(current_section)
        
        return sections
    
    def _is_section_header(self, text: str) -> bool:
        """
        Check if text looks like a section header using patterns.
        """
        text_lower = text.lower().strip()
        
        # Check numbered sections (1. Introduction, 2.1 Scope, etc.)
        if re.match(r'^\d+\.?\d*\.?\s+\w+', text):
            return True
        
        # Check against known section names
        for section_key in self.section_mapping.keys():
            if section_key in text_lower:
                return True
        
        # Check common header patterns
        header_patterns = [
            r'^(introduction|overview|scope|requirements|specifications?)$',
            r'^(functional|non-functional|user|system|api|ui|ux|workflow)',
            r'^(appendix|glossary|references|revision)',
        ]
        
        for pattern in header_patterns:
            if re.match(pattern, text_lower):
                return True
        
        return False
    
    def _map_to_confluence_page(self, section_title: str) -> str:
        """Map a section title to a Confluence page type."""
        title_lower = section_title.lower()
        
        # Remove numbering (1. Introduction -> introduction)
        title_clean = re.sub(r'^\d+\.?\d*\.?\s*', '', title_lower).strip()
        
        # Direct mapping
        if title_clean in self.section_mapping:
            return self.section_mapping[title_clean]
        
        # Partial matching
        for key, page_type in self.section_mapping.items():
            if key in title_clean or title_clean in key:
                return page_type
        
        # Default to Product Overview for unknown sections
        return "Product Overview"
    
    def _is_requirement(self, text: str) -> bool:
        """Check if text represents a requirement."""
        patterns = [
            r'^(REQ|FR|NFR|SR|UR|BR)[-_]?\d+',  # REQ-001, FR001, etc.
            r'^\d+\.\d+\.\d+',  # 1.2.3 format
            r'^(shall|must|should|will)\s',  # Requirement keywords
            r'^\[R\d+\]',  # [R1], [R2] format
        ]
        
        for pattern in patterns:
            if re.match(pattern, text, re.IGNORECASE):
                return True
        return False
    
    def _extract_metadata(self, doc: Document) -> Dict[str, Any]:
        """Extract document metadata."""
        metadata = {}
        
        try:
            core_props = doc.core_properties
            metadata["author"] = core_props.author
            metadata["created"] = str(core_props.created) if core_props.created else None
            metadata["modified"] = str(core_props.modified) if core_props.modified else None
            metadata["title"] = core_props.title
            metadata["subject"] = core_props.subject
        except Exception as e:
            logger.warning(f"Could not extract metadata: {e}")
        
        return metadata
    
    def generate_tasks(self, sections: List[SRSSection]) -> List[Dict[str, Any]]:
        """
        Generate initial tasks based on extracted sections.
        Assigns tasks to team members in round-robin fashion per section type.
        """
        tasks = []
        task_counter = 0
        
        # Track which section types we've seen
        seen_page_types = set()
        
        for section in sections:
            page_type = section.confluence_page
            
            # Only generate tasks once per page type
            if page_type in seen_page_types:
                continue
            seen_page_types.add(page_type)
            
            # Get task templates for this page type
            templates = TASK_TEMPLATES.get(page_type, [])
            
            for template in templates:
                assignee_idx = template["assignee_idx"] % len(self.team_members)
                assignee = self.team_members[assignee_idx]
                
                task = {
                    "title": f"{template['title']} - {section.title}",
                    "description": f"Task generated from SRS section: {section.title}\n\n"
                                   f"Related to: {page_type}\n"
                                   f"Section content summary available in Confluence.",
                    "assignee": assignee,
                    "section": section.title,
                    "confluence_page": page_type,
                    "priority": "Medium",
                    "labels": ["srs-generated", page_type.lower().replace(" ", "-").replace("/", "-")],
                }
                tasks.append(task)
                task_counter += 1
        
        logger.info(f"Generated {len(tasks)} tasks from {len(seen_page_types)} section types")
        return tasks
    
    def generate_user_stories(self, sections: List[SRSSection]) -> List[Dict[str, Any]]:
        """
        Generate user stories from functional requirements sections.
        """
        user_stories = []
        
        for section in sections:
            if section.confluence_page != "Feature Pages":
                continue
            
            # Parse requirements into user stories
            for req in section.requirements:
                story = self._requirement_to_user_story(req, section.title)
                if story:
                    user_stories.append(story)
            
            # Also try to extract from content
            content_stories = self._extract_stories_from_content(section.content)
            user_stories.extend(content_stories)
        
        logger.info(f"Generated {len(user_stories)} user stories")
        return user_stories
    
    def _requirement_to_user_story(self, requirement: str, section_title: str) -> Optional[Dict[str, Any]]:
        """Convert a requirement statement to a user story format."""
        # Try to parse "shall" statements into user stories
        shall_match = re.search(
            r'(the system|the application|the user|users?)\s+(shall|must|should|will)\s+(.+)',
            requirement,
            re.IGNORECASE
        )
        
        if shall_match:
            actor = shall_match.group(1)
            action = shall_match.group(3)
            
            # Determine user type
            if "user" in actor.lower():
                user_type = "user"
            else:
                user_type = "system"
            
            return {
                "title": f"As a {user_type}, {action[:80]}...",
                "description": requirement,
                "acceptance_criteria": [f"Verify: {requirement}"],
                "source_section": section_title,
                "story_points": 3,  # Default estimate
            }
        
        return None
    
    def _extract_stories_from_content(self, content: str) -> List[Dict[str, Any]]:
        """Extract user stories from section content."""
        stories = []
        
        # Look for "As a... I want... So that..." patterns
        pattern = r'As\s+a[n]?\s+(.+?),?\s+I\s+want\s+(.+?)(?:,?\s+so\s+that\s+(.+?))?(?:\.|$)'
        matches = re.findall(pattern, content, re.IGNORECASE | re.MULTILINE)
        
        for match in matches:
            user_type, want, benefit = match
            stories.append({
                "title": f"As a {user_type}, I want {want[:50]}...",
                "description": f"As a {user_type}, I want {want}" + 
                              (f", so that {benefit}" if benefit else ""),
                "acceptance_criteria": [],
                "source_section": "Extracted from content",
                "story_points": 3,
            })
        
        return stories


def get_srs_parser() -> SRSParser:
    """Factory function to get an SRS parser instance."""
    return SRSParser()


def build_confluence_page_html(
    page_type: str,
    section: SRSSection,
    project_name: str
) -> str:
    """
    Build HTML content for a Confluence page based on section type.
    """
    html_parts = []
    
    # Header
    html_parts.append(f"""
    <h1>{section.title}</h1>
    <p><strong>Project:</strong> {project_name}</p>
    <p><strong>Page Type:</strong> {page_type}</p>
    <hr/>
    """)
    
    # Content based on page type
    if page_type == "Product Overview":
        html_parts.append("""
        <h2>Overview</h2>
        <ac:structured-macro ac:name="info">
            <ac:rich-text-body>
                <p>This page contains the product overview and introduction from the SRS document.</p>
            </ac:rich-text-body>
        </ac:structured-macro>
        """)
    
    elif page_type == "Feature Pages":
        html_parts.append("""
        <h2>Functional Requirements</h2>
        <ac:structured-macro ac:name="note">
            <ac:rich-text-body>
                <p>Features and requirements extracted from the SRS. User stories are linked below.</p>
            </ac:rich-text-body>
        </ac:structured-macro>
        """)
    
    elif page_type == "NFR":
        html_parts.append("""
        <h2>Non-Functional Requirements</h2>
        <table>
            <tr>
                <th>Category</th>
                <th>Requirement</th>
                <th>Target</th>
            </tr>
        """)
    
    # Main content
    content_html = section.content.replace('\n', '<br/>')
    html_parts.append(f"""
    <h2>Details</h2>
    <div class="srs-content">
        {content_html}
    </div>
    """)
    
    # Requirements list if any
    if section.requirements:
        html_parts.append("<h2>Requirements</h2><ul>")
        for req in section.requirements:
            html_parts.append(f"<li>{req}</li>")
        html_parts.append("</ul>")
    
    # Footer with metadata
    html_parts.append(f"""
    <hr/>
    <p><em>Auto-generated from SRS document. Last updated: {{date}}</em></p>
    """)
    
    return "\n".join(html_parts)
